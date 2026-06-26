from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CN_BODY_FONT = "宋体"
CN_HEADING_FONT = "黑体"
EN_FONT = "Times New Roman"
ACCENT = RGBColor(31, 78, 121)


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None,
                 cn_font: str = CN_BODY_FONT, en_font: str = EN_FONT) -> None:
    run.font.name = en_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size_pt: float, cn_font: str, en_font: str = EN_FONT,
                   bold: bool = False, color: RGBColor | None = None) -> None:
    style.font.name = en_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_paragraph_format(paragraph, line_spacing: float = 1.5,
                         first_line_chars: float | None = 2,
                         before_pt: float = 0, after_pt: float = 6) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if first_line_chars is not None:
        fmt.first_line_indent = Pt(10.5 * first_line_chars)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.7)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5, CN_BODY_FONT)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Title"], 18, CN_HEADING_FONT, bold=True, color=ACCENT)
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Title"].paragraph_format.space_after = Pt(18)

    set_style_font(styles["Heading 1"], 15, CN_HEADING_FONT, bold=True, color=ACCENT)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], 13, CN_HEADING_FONT, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True


def add_mixed_paragraph(doc: Document, parts: list[dict], style: str | None = None,
                        indent: bool = True) -> None:
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_format(paragraph, first_line_chars=2 if indent else None)
    for part in parts:
        run = paragraph.add_run(part["text"])
        set_run_font(
            run,
            size_pt=part.get("size", 10.5),
            bold=part.get("bold"),
            italic=part.get("italic"),
            color=ACCENT if part.get("accent") else None,
            cn_font=part.get("cn_font", CN_BODY_FONT),
        )


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, 180, 180, 180, 180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=10.5, bold=True, cn_font=CN_BODY_FONT, color=ACCENT)
    doc.add_paragraph()


def add_matrix_table(doc: Document, rows: list[dict]) -> None:
    headers = ["主题", "理论基础", "方法/数据", "核心结论", "可借鉴点"]
    keys = ["topic", "theory", "method", "finding", "use"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_repeat_table_header(table.rows[0])

    widths = [Cm(2.6), Cm(3.2), Cm(3.2), Cm(4.3), Cm(3.8)]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size_pt=9.5, bold=True, cn_font=CN_HEADING_FONT, color=ACCENT)

    for item in rows:
        cells = table.add_row().cells
        for i, key in enumerate(keys):
            cell = cells[i]
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(item[key])
            set_run_font(run, size_pt=9, cn_font=CN_BODY_FONT)


def add_references(doc: Document, references: list[str]) -> None:
    for ref in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(21)
        paragraph.paragraph_format.first_line_indent = Pt(-21)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(ref)
        set_run_font(run, size_pt=10, cn_font=CN_BODY_FONT)


def build_doc(content: dict, output: Path) -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title_run = title.add_run(content["title"])
    set_run_font(title_run, size_pt=18, bold=True, cn_font=CN_HEADING_FONT, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    sub_run = subtitle.add_run(content["subtitle"])
    set_run_font(sub_run, size_pt=10.5, cn_font=CN_BODY_FONT, color=RGBColor(89, 89, 89))

    add_callout(doc, content["summary"])

    for section in content["sections"]:
        doc.add_heading(section["heading"], level=1)
        for paragraph in section["paragraphs"]:
            add_mixed_paragraph(doc, paragraph)

    doc.add_heading("文献矩阵示例", level=1)
    add_matrix_table(doc, content["matrix"])

    doc.add_heading("参考文献示例", level=1)
    add_references(doc, content["references"])

    doc.add_section(WD_SECTION_START.CONTINUOUS)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def validate_docx(output: Path, expected_text: list[str]) -> None:
    with zipfile.ZipFile(output) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    missing = [text for text in expected_text if text not in xml]
    if missing:
        raise RuntimeError(f"生成后的 DOCX 缺少预期中文文本: {missing}")
    if "\ufffd" in xml or re.search(r"\?{4,}", xml):
        raise RuntimeError("生成后的 DOCX 出现替换字符或连续问号，疑似编码已经损坏。")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a UTF-8 safe Chinese Word document.")
    parser.add_argument("--input", default="../assets/sample_content.json", help="UTF-8 JSON content file.")
    parser.add_argument("--output", default="中文Word生成测试_无乱码.docx", help="Output DOCX path.")
    args = parser.parse_args()

    base_dir = Path(__file__).resolve().parent
    input_path = Path(args.input)
    if not input_path.is_absolute():
        input_path = base_dir / input_path
    output_path = Path(args.output)
    if not output_path.is_absolute():
        output_path = base_dir / output_path

    content = json.loads(input_path.read_text(encoding="utf-8-sig"))
    build_doc(content, output_path)
    validate_docx(output_path, [content["title"], "中文文件名", "AI washing", "文献矩阵示例"])
    print(f"OK: {output_path}")


if __name__ == "__main__":
    main()
