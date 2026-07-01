from __future__ import annotations

import argparse
import re
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from fix_equation_number_tabs import patch_docx as repair_equation_number_tabs

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

CN_PUNCT_MAP = str.maketrans({
    ",": "，",
    ";": "；",
    ":": "：",
    "?": "？",
    "!": "！",
})

BRACED_SUBSCRIPT_RE = re.compile(r"([A-Za-zΑ-Ωα-ωΣμλβεω]+)_\{([^{}]+)\}")
SIMPLE_SUBSCRIPT_RE = re.compile(r"([A-Za-zΑ-Ωα-ωΣμλβεω]+)_([A-Za-z0-9]+)")
EQUATION_NO_ONLY_RE = re.compile(r"^[\t\s（）()0-9]+$")

INLINE_SUBSCRIPT_AUDIT_PATTERNS = {
    "Resilienceit": "it",
    "AIWit": "it",
    "Controlsit": "it",
    "Channelit": "it",
    "Outcomeit": "it",
    "AnalystAttentionit": "it",
    "AIDisclosureit": "it",
    "AIPatentit": "it",
    "Returnim": "im",
    "PRkt": "kt",
    "Kit": "it",
    "Nikt": "ikt",
    "Ljt": "jt",
    "μi": "i",
    "λt": "t",
    "εit": "it",
    "νit": "it",
    "ξit": "it",
    "ωit": "it",
    "w1": "1",
    "w2": "2",
    "α1": "1",
    "α2": "2",
    "βk": "k",
    "δk": "k",
    "φk": "k",
    "κk": "k",
    "γ0": "0",
    "γ1": "1",
    "γ2": "2",
    "θ0": "0",
    "θ1": "1",
    "θ2": "2",
    "θ3": "3",
    "ρ0": "0",
    "ρ1": "1",
    "ρ2": "2",
    "ρ3": "3",
    "ρ4": "4",
    "ρ5": "5",
}


def is_caption_text(text: str) -> bool:
    text = text.strip()
    if re.match(r"^(图|表)\s*[0-9]+[A-Za-z]?\s*\S+", text):
        return not re.match(r"^(图|表)\s*[0-9]+[A-Za-z]?\s*(报告|显示|列示|给出|展示|说明|中|可见|表明)", text)
    if re.match(r"^(Figure|Table)\s+[0-9]+[A-Za-z]?\.", text):
        return True
    return False


def normalize_cn_text(text: str) -> str:
    if not text:
        return text
    # Preserve citekeys and English phrases as much as possible; only clean
    # spacing and punctuation where adjacent Chinese text makes the intent clear.
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[，。！？；：、])", "", text)
    text = re.sub(r"(?<=[（《“])\s+", "", text)
    text = re.sub(r"\s+(?=[）》”])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[0-9％%])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[0-9])", "", text)
    text = re.sub(r"(?<=[0-9])\s+(?=[至到—-])", "", text)
    text = re.sub(r"(?<=[至到—-])\s+(?=[0-9])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+([,;:?!])", r"\1", text)
    text = re.sub(r"([,;:?!])(?=[\u4e00-\u9fff])", lambda m: m.group(1).translate(CN_PUNCT_MAP), text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])([,;:?!])", lambda m: m.group(1).translate(CN_PUNCT_MAP), text)
    return text


def normalize_caption_text(text: str) -> str:
    text = normalize_cn_text(text)
    text = re.sub(r"^(图|表)\s*([0-9]+[A-Za-z]?)\s*", r"\1 \2 ", text)
    return text


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float | None = None, bold=None, italic=None, subscript=False):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.subscript = subscript


def set_style_fonts(doc: Document, lang: str):
    if lang == "cn":
        mappings = {
            "Normal": ("宋体", "Times New Roman", 10.5, False),
            "Body Text": ("宋体", "Times New Roman", 10.5, False),
            "First Paragraph": ("宋体", "Times New Roman", 10.5, False),
            "Heading 1": ("黑体", "Times New Roman", 15, True),
            "Heading 2": ("黑体", "Times New Roman", 13, True),
            "Heading 3": ("黑体", "Times New Roman", 12, True),
            "Caption": ("宋体", "Times New Roman", 10.5, False),
        }
    else:
        mappings = {
            "Normal": ("Times New Roman", "Times New Roman", 11, False),
            "Body Text": ("Times New Roman", "Times New Roman", 11, False),
            "First Paragraph": ("Times New Roman", "Times New Roman", 11, False),
            "Heading 1": ("Times New Roman", "Times New Roman", 15, True),
            "Heading 2": ("Times New Roman", "Times New Roman", 13, True),
            "Heading 3": ("Times New Roman", "Times New Roman", 12, True),
            "Caption": ("Times New Roman", "Times New Roman", 10.5, False),
        }

    for style_name, (east_asia, latin, size_pt, bold) in mappings.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = latin
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.font.size = Pt(size_pt)
        style.font.bold = bold


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def append_run(paragraph, text: str, east_asia: str, latin: str, size_pt: float, *, bold=None, italic=None, subscript=False):
    run = paragraph.add_run(text)
    set_run_fonts(run, east_asia, latin, size_pt, bold=bold, italic=italic, subscript=subscript)
    return run


def append_symbol(paragraph, base: str, sub: str | None, east_asia: str, latin: str, size_pt: float, *, italic=True, greek=False):
    append_run(paragraph, base, east_asia, latin, size_pt, italic=(False if greek else italic))
    if sub:
        append_run(paragraph, sub, east_asia, latin, size_pt - 1, italic=(False if greek else italic), subscript=True)


def reset_paragraph_runs(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def paragraph_has_fields(paragraph) -> bool:
    p = paragraph._element
    return bool(
        p.findall(f".//{{{W_NS}}}fldChar")
        or p.findall(f".//{{{W_NS}}}instrText")
        or p.findall(f".//{{{W_NS}}}hyperlink")
    )


def _find_next_inline_symbol(text: str, start: int):
    braced = BRACED_SUBSCRIPT_RE.search(text, start)
    simple = SIMPLE_SUBSCRIPT_RE.search(text, start)
    candidates = [m for m in (braced, simple) if m]
    if not candidates:
        return None
    match = min(candidates, key=lambda m: (m.start(), m.end()))
    if match.re is SIMPLE_SUBSCRIPT_RE and braced and braced.start() == match.start():
        return braced
    return match


def rebuild_inline_symbol_paragraph(paragraph, lang: str) -> bool:
    if paragraph_has_fields(paragraph) or paragraph_has_drawing(paragraph):
        return False
    if paragraph._element.findall(f".//{{{M_NS}}}oMath") or paragraph._element.findall(f".//{{{M_NS}}}oMathPara"):
        return False

    original = paragraph.text
    if "_{" not in original and not SIMPLE_SUBSCRIPT_RE.search(original):
        return False

    text = original.replace("`", "")
    east_asia = "宋体" if lang == "cn" else "Times New Roman"
    latin = "Times New Roman"
    size_pt = 10.5 if lang == "cn" else 11

    matches = list(BRACED_SUBSCRIPT_RE.finditer(text))
    if not matches and not SIMPLE_SUBSCRIPT_RE.search(text):
        return False

    clear_paragraph(paragraph)
    idx = 0
    while idx < len(text):
        match = _find_next_inline_symbol(text, idx)
        if not match:
            append_run(paragraph, text[idx:], east_asia, latin, size_pt)
            break
        if match.start() > idx:
            append_run(paragraph, text[idx:match.start()], east_asia, latin, size_pt)
        base, sub = match.group(1), match.group(2)
        append_symbol(paragraph, base, sub, east_asia, latin, size_pt, italic=False, greek=(base in {"μ", "λ", "ε", "α", "β", "ω", "Σ"}))
        idx = match.end()
    return True


def rebuild_collapsed_symbol_paragraph(paragraph, lang: str) -> bool:
    text = paragraph.text.strip()
    east_asia = "宋体" if lang == "cn" else "Times New Roman"
    latin = "Times New Roman"
    size_pt = 10.5 if lang == "cn" else 11

    symbol_map = [
        ("Resilienceit", ("Resilience", "it")),
        ("Resilienceit²", ("Resilience", "it")),
        ("AIWit", ("AIW", "it")),
        ("AIWit²", ("AIW", "it")),
        ("Controlsit", ("Controls", "it")),
        ("Channelit", ("Channel", "it")),
        ("Outcomeit", ("Outcome", "it")),
        ("AnalystAttentionit", ("AnalystAttention", "it")),
        ("AIDisclosureit", ("AIDisclosure", "it")),
        ("AIPatentit", ("AIPatent", "it")),
        ("Returnim", ("Return", "im")),
        ("PRkt", ("PR", "kt")),
        ("Kit", ("K", "it")),
        ("Nikt", ("N", "ikt")),
        ("μi", ("μ", "i")),
        ("λt", ("λ", "t")),
        ("εit", ("ε", "it")),
        ("νit", ("ν", "it")),
        ("ξit", ("ξ", "it")),
        ("ωit", ("ω", "it")),
        ("w1", ("w", "1")),
        ("w2", ("w", "2")),
        ("α1", ("α", "1")),
        ("α2", ("α", "2")),
        ("βk", ("β", "k")),
        ("δk", ("δ", "k")),
        ("φk", ("φ", "k")),
        ("κk", ("κ", "k")),
        ("γ0", ("γ", "0")),
        ("γ1", ("γ", "1")),
        ("γ2", ("γ", "2")),
        ("θ0", ("θ", "0")),
        ("θ1", ("θ", "1")),
        ("θ2", ("θ", "2")),
        ("θ3", ("θ", "3")),
        ("ρ0", ("ρ", "0")),
        ("ρ1", ("ρ", "1")),
        ("ρ2", ("ρ", "2")),
        ("ρ3", ("ρ", "3")),
        ("ρ4", ("ρ", "4")),
        ("ρ5", ("ρ", "5")),
    ]
    if not any(token in text for token, _ in symbol_map):
        return False
    if paragraph_has_fields(paragraph) or paragraph_has_drawing(paragraph):
        return False

    reset_paragraph_runs(paragraph)
    cursor = 0
    while cursor < len(text):
        next_match = None
        next_idx = len(text)
        for token, payload in symbol_map:
            idx = text.find(token, cursor)
            if idx != -1 and idx < next_idx:
                next_match = (idx, token, payload)
                next_idx = idx
        if next_match is None:
            append_run(paragraph, text[cursor:], east_asia, latin, size_pt)
            break
        idx, token, (base, sub) = next_match
        if idx > cursor:
            append_run(paragraph, text[cursor:idx], east_asia, latin, size_pt)
        append_symbol(paragraph, base, sub, east_asia, latin, size_pt, italic=False, greek=(base in {"μ", "λ", "ε", "ν", "ξ", "ω", "α", "β", "γ", "δ", "φ", "κ", "ρ", "θ"}))
        cursor = idx + len(token)
    return True


def _m_el(tag: str):
    return OxmlElement(f"m:{tag}")


def _w_el(tag: str):
    return OxmlElement(f"w:{tag}")


def math_text(text: str):
    run = _m_el("r")
    r_pr = _m_el("rPr")
    r_pr.append(_m_el("nor"))
    run.append(r_pr)
    t = _m_el("t")
    t.text = text
    run.append(t)
    return run


def math_rich_text(text: str, *, italic: bool = False):
    run = _m_el("r")
    r_pr = _m_el("rPr")
    if not italic:
        r_pr.append(_m_el("nor"))
    run.append(r_pr)
    t = _m_el("t")
    t.text = text
    run.append(t)
    return run


def math_sub(base: str, sub: str, *, italic_base: bool = False):
    node = _m_el("sSub")
    e = _m_el("e")
    e.append(math_rich_text(base, italic=italic_base))
    sub_el = _m_el("sub")
    sub_el.append(math_rich_text(sub, italic=False))
    node.append(e)
    node.append(sub_el)
    return node


def math_sub_expr(expr, sub: str):
    node = _m_el("sSub")
    e = _m_el("e")
    e.append(expr)
    sub_el = _m_el("sub")
    sub_el.append(math_rich_text(sub, italic=False))
    node.append(e)
    node.append(sub_el)
    return node


def math_subsup(base: str, sub: str, sup: str, *, italic_base: bool = False):
    node = _m_el("sSubSup")
    e = _m_el("e")
    e.append(math_rich_text(base, italic=italic_base))
    sub_el = _m_el("sub")
    sub_el.append(math_rich_text(sub, italic=False))
    sup_el = _m_el("sup")
    sup_el.append(math_rich_text(sup, italic=False))
    node.append(e)
    node.append(sub_el)
    node.append(sup_el)
    return node


def math_group(elements):
    group = _m_el("d")
    d_pr = _m_el("dPr")
    beg = _m_el("begChr")
    beg.set(qn("m:val"), "(")
    end = _m_el("endChr")
    end.set(qn("m:val"), ")")
    d_pr.append(beg)
    d_pr.append(end)
    group.append(d_pr)
    e = _m_el("e")
    for element in elements:
        e.append(element)
    group.append(e)
    return group


def math_frac(num_elements, den_elements):
    frac = _m_el("f")
    num = _m_el("num")
    den = _m_el("den")
    for element in num_elements:
        num.append(element)
    for element in den_elements:
        den.append(element)
    frac.append(num)
    frac.append(den)
    return frac


def math_overbar(element):
    acc = _m_el("acc")
    acc_pr = _m_el("accPr")
    chr_el = _m_el("chr")
    chr_el.set(qn("m:val"), "¯")
    acc_pr.append(chr_el)
    acc.append(acc_pr)
    e = _m_el("e")
    e.append(element)
    acc.append(e)
    return acc


def _math_component_text(component) -> str:
    return "".join(t.text or "" for t in component.findall(f".//{{{M_NS}}}t"))


def _auto_wrap_math_lines(lines, *, max_chars: int = 78, soft_threshold: int = 52):
    wrapped = []
    for line in lines:
        line_text = "".join(_math_component_text(comp) for comp in line)
        if len(line_text) <= max_chars:
            wrapped.append(line)
            continue

        current = []
        current_chars = 0
        for comp in line:
            comp_text = _math_component_text(comp)
            stripped = comp_text.strip()
            should_break = (
                current
                and current_chars >= soft_threshold
                and (
                    stripped.startswith("+")
                    or stripped.startswith("-")
                    or stripped.startswith("×")
                )
            )
            if should_break:
                wrapped.append(current)
                current = [deepcopy(comp)]
                current_chars = len(comp_text)
                continue
            current.append(deepcopy(comp))
            current_chars += len(comp_text)
        if current:
            wrapped.append(current)
    return wrapped


def build_omml_multiline_paragraph(lines):
    o_math_para = _m_el("oMathPara")
    o_math_para_pr = _m_el("oMathParaPr")
    jc_math = _m_el("jc")
    jc_math.set(qn("m:val"), "center")
    o_math_para_pr.append(jc_math)
    o_math_para.append(o_math_para_pr)
    o_math = _m_el("oMath")
    eq_arr = _m_el("eqArr")
    eq_arr_pr = _m_el("eqArrPr")
    for tag, value in (("maxDist", "1"), ("objDist", "0"), ("rSp", "1")):
        node = _m_el(tag)
        node.set(qn("m:val"), value)
        eq_arr_pr.append(node)
    eq_arr.append(eq_arr_pr)
    for line in lines:
        e = _m_el("e")
        for comp in line:
            e.append(comp)
        eq_arr.append(e)
    o_math.append(eq_arr)
    o_math_para.append(o_math)
    p_pr = _w_el("pPr")
    jc = _w_el("jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    para = OxmlElement("w:p")
    para.append(p_pr)
    para.append(o_math_para)
    return para


def build_equation_numbered_paragraph(lines, eq_no: int, *, lang: str):
    lines = _auto_wrap_math_lines(lines)
    para = OxmlElement("w:p")
    p_pr = _w_el("pPr")
    jc = _w_el("jc")
    jc.set(qn("w:val"), "left")
    p_pr.append(jc)
    ind = _w_el("ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    p_pr.append(ind)
    spacing = _w_el("spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    tabs = _w_el("tabs")
    center_tab = _w_el("tab")
    center_tab.set(qn("w:val"), "center")
    center_tab.set(qn("w:pos"), "4500")
    tabs.append(center_tab)
    tab = _w_el("tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    p_pr.append(tabs)
    para.append(p_pr)

    if len(lines) > 1:
        leading_tab_run = _w_el("r")
        leading_tab_run.append(_w_el("tab"))
        para.append(leading_tab_run)

        o_math_para = _m_el("oMathPara")
        o_math_para_pr = _m_el("oMathParaPr")
        jc_math = _m_el("jc")
        jc_math.set(qn("m:val"), "center")
        o_math_para_pr.append(jc_math)
        o_math_para.append(o_math_para_pr)

        o_math = _m_el("oMath")
        eq_arr = _m_el("eqArr")
        eq_arr_pr = _m_el("eqArrPr")
        for tag, value in (("maxDist", "1"), ("objDist", "0"), ("rSp", "1")):
            node = _m_el(tag)
            node.set(qn("m:val"), value)
            eq_arr_pr.append(node)
        eq_arr.append(eq_arr_pr)
        for line in lines:
            e = _m_el("e")
            for comp in line:
                e.append(deepcopy(comp))
            eq_arr.append(e)
        o_math.append(eq_arr)
        o_math_para.append(o_math)
        para.append(o_math_para)
    else:
        for line in lines:
            leading_tab_run = _w_el("r")
            leading_tab_run.append(_w_el("tab"))
            para.append(leading_tab_run)

            o_math = _m_el("oMath")
            for comp in line:
                o_math.append(deepcopy(comp))
            para.append(o_math)

    tab_run = _w_el("r")
    tab_run.append(_w_el("tab"))
    para.append(tab_run)

    num_run = _w_el("r")
    r_pr = _w_el("rPr")
    r_fonts = _w_el("rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体" if lang == "cn" else "Times New Roman")
    r_pr.append(r_fonts)
    sz = _w_el("sz")
    sz.set(qn("w:val"), "21" if lang == "cn" else "22")
    r_pr.append(sz)
    num_run.append(r_pr)
    t = _w_el("t")
    t.text = f"（{eq_no}）" if lang == "cn" else f"({eq_no})"
    num_run.append(t)
    para.append(num_run)
    return para


def _clone_math_line_from_eqarr(eqarr_line):
    cloned = []
    for comp in list(eqarr_line):
        cloned.append(deepcopy(comp))
    return cloned


def normalize_existing_omml_numbered_paragraphs(doc: Document, lang: str):
    body = doc._element.body
    for p in list(body):
        if p.tag != qn("w:p"):
            continue
        o_math_para = p.find(qn("m:oMathPara"))
        if o_math_para is None:
            continue
        texts = [t.text or "" for t in p.findall(f".//{{{W_NS}}}t")]
        joined = "".join(texts)
        eq_no = _parse_equation_number(joined)
        if eq_no is None:
            continue
        eqarr = o_math_para.find(f".//{{{M_NS}}}eqArr")
        if eqarr is None:
            continue
        lines = [_clone_math_line_from_eqarr(e) for e in eqarr.findall(f"./{{{M_NS}}}e")]
        if not lines:
            continue
        new_p = build_equation_numbered_paragraph(lines, eq_no, lang=lang)
        body.insert(body.index(p), new_p)
        body.remove(p)


def _extract_math_lines_from_omath(omath):
    eqarr = omath.find(qn("m:eqArr"))
    if eqarr is not None:
        lines = []
        for e in eqarr.findall(qn("m:e")):
            lines.append(_clone_math_line_from_eqarr(e))
        return lines
    return [[deepcopy(comp) for comp in list(omath)]]


def _extract_math_lines_from_paragraph_xml(p):
    o_math_para = p.find(qn("m:oMathPara"))
    if o_math_para is not None:
        o_math = o_math_para.find(qn("m:oMath"))
        if o_math is not None:
            return _extract_math_lines_from_omath(o_math)
    o_math = p.find(qn("m:oMath"))
    if o_math is not None:
        return _extract_math_lines_from_omath(o_math)
    return None


def _extract_math_lines_from_table_xml(tbl):
    for cell in tbl.findall(f".//{{{W_NS}}}tc"):
        for p in cell.findall(f"./{{{W_NS}}}p"):
            lines = _extract_math_lines_from_paragraph_xml(p)
            if lines:
                return lines
    return None


def _extract_equation_number_from_table_xml(tbl):
    for cell in reversed(tbl.findall(f".//{{{W_NS}}}tc")):
        texts = "".join(t.text or "" for t in cell.findall(f".//{{{W_NS}}}t")).strip()
        eq_no = _parse_equation_number(texts)
        if eq_no is not None:
            return eq_no
    return None


def table_visible_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def normalize_display_equation_paragraphs(doc: Document, lang: str):
    body = doc._element.body
    eq_no = 1
    for node in list(body):
        lines = None
        if node.tag == qn("w:p"):
            lines = _extract_math_lines_from_paragraph_xml(node)
            if lines:
                visible_text = _paragraph_text_xml(node).strip()
                if visible_text and _parse_equation_number(visible_text) is None:
                    lines = None
        elif node.tag == qn("w:tbl"):
            lines = _extract_math_lines_from_table_xml(node)
        if not lines:
            continue
        new_p = build_equation_numbered_paragraph(lines, eq_no, lang=lang)
        body.insert(body.index(node), new_p)
        body.remove(node)
        eq_no += 1


def build_text_paragraph(text: str, align: str = "right"):
    para = OxmlElement("w:p")
    p_pr = _w_el("pPr")
    jc = _w_el("jc")
    jc.set(qn("w:val"), align)
    p_pr.append(jc)
    para.append(p_pr)
    run = _w_el("r")
    r_pr = _w_el("rPr")
    r_fonts = _w_el("rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_pr.append(r_fonts)
    sz = _w_el("sz")
    sz.set(qn("w:val"), "21")
    r_pr.append(sz)
    run.append(r_pr)
    t = _w_el("t")
    t.text = text
    run.append(t)
    para.append(run)
    return para


def _set_cell_width(tc, width_twips: int):
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = _w_el("tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(tc, margin_twips: int = 0):
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = _w_el("tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = _w_el(side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def _set_table_no_borders(tbl):
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = _w_el("tblPr")
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = _w_el("tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = _w_el(edge)
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _set_table_full_width(tbl):
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = _w_el("tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = _w_el("tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = _w_el("jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)


def _set_paragraph_spacing_xml(p, *, before: int = 80, after: int = 80, align: str = "center"):
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = _w_el("pPr")
        p.insert(0, p_pr)
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = _w_el("jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), align)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = _w_el("spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:lineRule"), "auto")
    if qn("w:line") in spacing.attrib:
        del spacing.attrib[qn("w:line")]


def _force_table_cell_paragraph_xml(p, *, align: str = "center"):
    """Make table-cell paragraph geometry explicit at OOXML level.

    Word/WPS may inherit first-line character indentation from body styles even
    when python-docx point indents are zero.  The char-based attributes are the
    important guardrail for Chinese manuscripts.
    """
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = _w_el("pPr")
        p.insert(0, p_pr)

    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = _w_el("ind")
        p_pr.append(ind)
    for attr in ("left", "right", "firstLine", "firstLineChars"):
        ind.set(qn(f"w:{attr}"), "0")
    for attr in ("hanging", "hangingChars"):
        ind.attrib.pop(qn(f"w:{attr}"), None)

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = _w_el("jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), align)


def force_table_cell_paragraph_geometry(paragraph, *, lang: str, align: str = "center"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(16 if lang == "cn" else 14)
    _force_table_cell_paragraph_xml(paragraph._element, align=align)


def force_cell_vertical_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = _w_el("vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")


def build_equation_table(equation_p, eq_no: int, *, lang: str = "cn"):
    """Build a borderless two-cell equation block with same-line numbering."""
    tbl = _w_el("tbl")
    tbl_pr = _w_el("tblPr")
    tbl_w = _w_el("tblW")
    tbl_w.set(qn("w:w"), "9000")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    jc = _w_el("jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)
    tbl.append(tbl_pr)
    _set_table_no_borders(tbl)

    grid = _w_el("tblGrid")
    for width in ("8000", "1000"):
        col = _w_el("gridCol")
        col.set(qn("w:w"), width)
        grid.append(col)
    tbl.append(grid)

    tr = _w_el("tr")
    for width in (8000, 1000):
        tc = _w_el("tc")
        _set_cell_width(tc, width)
        _set_cell_margins(tc, 0)
        v_align = _w_el("vAlign")
        v_align.set(qn("w:val"), "center")
        tc.get_or_add_tcPr().append(v_align)
        tr.append(tc)
    tbl.append(tr)

    eq_cell, no_cell = tr.findall(qn("w:tc"))
    eq_p = deepcopy(equation_p)
    _set_paragraph_spacing_xml(eq_p, before=80, after=80, align="center")
    eq_cell.append(eq_p)

    no_p = build_text_paragraph(f"（{eq_no}）" if lang == "cn" else f"({eq_no})", align="right")
    _set_paragraph_spacing_xml(no_p, before=80, after=80, align="right")
    no_cell.append(no_p)
    return tbl


def replace_with_omml_block(paragraph, lines, eq_no: int | None = None, *, lang: str = "cn"):
    parent = paragraph._element.getparent()
    new_p = build_equation_numbered_paragraph(lines, eq_no, lang=lang) if eq_no is not None else build_omml_multiline_paragraph(lines)
    paragraph._element.addnext(new_p)
    parent.remove(paragraph._element)


def set_paragraph_format(paragraph, lang: str, in_table: bool = False):
    pf = paragraph.paragraph_format
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if in_table:
        force_table_cell_paragraph_geometry(paragraph, lang=lang, align="center")
        return

    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""
    has_drawing = paragraph_has_drawing(paragraph)
    if has_drawing:
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = None
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if style_name.startswith("Heading"):
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20 if lang == "cn" else 16)
    elif is_caption_text(text):
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18 if lang == "cn" else 14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.first_line_indent = Cm(0.74) if lang == "cn" else Pt(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY if lang == "cn" else WD_LINE_SPACING.ONE_POINT_FIVE
        pf.line_spacing = Pt(20 if lang == "cn" else 18)


def set_table_borders(cell, *, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    def apply(edge: str, spec):
        node = borders.find(qn(f"w:{edge}"))
        if spec is None:
            if node is not None:
                borders.remove(node)
            return
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        for k, v in spec.items():
            node.set(qn(f"w:{k}"), str(v))

    apply("top", top)
    apply("bottom", bottom)
    apply("left", left)
    apply("right", right)


def is_equation_table(table) -> bool:
    # python-docx's namespace-aware findall can be brittle for inserted OMML.
    # The serialized table XML is a safer marker for equation layout tables.
    xml = table._tbl.xml
    return "<m:oMath" in xml or "<m:oMathPara" in xml


def normalize_equation_table(table, lang: str):
    _set_table_no_borders(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for tr_height in list(table._tbl.iter(qn("w:trHeight"))):
        parent = tr_height.getparent()
        if parent is not None:
            parent.remove(tr_height)
    for p_xml in table._tbl.iter(qn("w:p")):
        _set_paragraph_spacing_xml(p_xml, before=160, after=160, align="center")
        spacing = p_xml.find(qn("w:pPr")).find(qn("w:spacing"))
        spacing.set(qn("w:lineRule"), "auto")
        spacing.attrib.pop(qn("w:line"), None)
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_table_borders(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
                if _is_equation_paragraph_xml(p._element):
                    _set_paragraph_spacing_xml(p._element, before=140, after=140, align="center")
                else:
                    _set_paragraph_spacing_xml(p._element, before=140, after=140, align="right")


def format_tables(doc: Document, lang: str):
    top_rule = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_rule = {"val": "single", "sz": "4", "space": "0", "color": "000000"}
    bottom_rule = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    none_rule = {"val": "nil"}
    east_asia = "宋体" if lang == "cn" else "Times New Roman"
    latin = "Times New Roman"
    body_size = 10.5 if lang == "cn" else 10.5

    for table in doc.tables:
        if is_equation_table(table):
            normalize_equation_table(table, lang)
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _set_table_full_width(table._tbl)
        last_row = len(table.rows) - 1
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                force_cell_vertical_center(cell)
                for p in cell.paragraphs:
                    set_paragraph_format(p, lang, in_table=True)
                    force_table_cell_paragraph_geometry(p, lang=lang, align="center")
                    for run in p.runs:
                        set_run_fonts(run, east_asia, latin, body_size, bold=(r_idx == 0))
                set_table_borders(
                    cell,
                    top=top_rule if r_idx == 0 else none_rule,
                    bottom=mid_rule if r_idx == 0 else (bottom_rule if r_idx == last_row else none_rule),
                    left=none_rule,
                    right=none_rule,
                )


def normalize_equation_source(text: str) -> str:
    text = text.strip().strip("`")
    text = text.replace("\\\\", "\\")
    text = text.replace("\\left", "").replace("\\right", "")
    text = text.replace("\\times", "×")
    text = text.replace("\\sum", "sum")
    text = re.sub(r"\s+", " ", text)
    return text


def _has_formula_token(text: str, *tokens: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return all(token in compact for token in tokens)


def rebuild_equation_paragraph(paragraph, lang: str):
    text = paragraph.text.strip()
    normalized = normalize_equation_source(text)
    if _has_formula_token(normalized, "Stability_{it}=", "SD(Return_{im})"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Stability", "it"),
                math_text(" = SD"),
                math_group([math_sub("Return", "im")]),
            ]],
            eq_no=1,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "Growth_{it}=", "Revenue_{it}-Revenue_{i,t-3}"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Growth", "it"),
                math_text(" = "),
                math_sub("Revenue", "it"),
                math_text(" - "),
                math_sub("Revenue", "i,t-3"),
            ]],
            eq_no=2,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "CR_{it}=", "z(-Stability_{it})", "z(Growth_{it})"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("CR", "it"),
                math_text(" = "),
                math_sub("w", "1"),
                math_text(" × z"),
                math_group([math_text("-"), math_sub("Stability", "it")]),
                math_text(" + "),
                math_sub("w", "2"),
                math_text(" × z"),
                math_group([math_sub("Growth", "it")]),
            ]],
            eq_no=3,
            lang=lang,
        )
        return True
    if text.startswith("W_kl = sum_p I(k in p) I(l in p), k != l"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("W", "kl"),
                math_text(" = "),
                math_sub("Σ", "p"),
                math_text(" I"),
                math_group([math_text("k in p")]),
                math_text(" I"),
                math_group([math_text("l in p")]),
                math_text(", k ≠ l"),
            ]],
            eq_no=1,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "PR_{kt}=", "sum_{j\\inM(k)}", "PR_{jt}/L_{jt}"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("PR", "kt"),
                math_text(" = (1-d) + d "),
                math_sub("Σ", "j∈M(k)"),
                math_text(" "),
                math_sub("PR", "jt"),
                math_text("/"),
                math_sub("L", "jt"),
            ]],
            eq_no=1,
            lang=lang,
        )
        return True
    if text.startswith("PR_k = (1-d)/N + d sum_l W_lk PR_l / sum_m W_lm"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("PR", "k"),
                    math_text(" = "),
                    math_group([math_text("1-d")]),
                    math_text("/N + d "),
                    math_sub("Σ", "l"),
                    math_text(" "),
                    math_sub("W", "lk"),
                    math_text(" "),
                    math_sub("PR", "l"),
                    math_text(" / "),
                    math_sub("Σ", "m"),
                    math_text(" "),
                    math_sub("W", "lm"),
                ],
            ],
            eq_no=2,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "AIPatent_{it}=", "sum_{k\\inK_{it}}", "PR_{kt}", "N_{ikt}"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("AIPatent", "it"),
                math_text(" = "),
                math_sub("Σ", "k∈K_it"),
                math_text(" "),
                math_sub("PR", "kt"),
                math_text(" × "),
                math_sub("N", "ikt"),
            ]],
            eq_no=2,
            lang=lang,
        )
        return True
    if text.startswith("AIPatent_it = sum_k N_ikt PR_kt"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("AIPatent", "it"),
                math_text(" = "),
                math_sub("Σ", "k"),
                math_text(" "),
                math_sub("N", "ikt"),
                math_text(" "),
                math_sub("PR", "kt"),
            ]],
            eq_no=3,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "AIW_{it}=", "\\overline{AIDisclosure}_{t}", "\\sigma(AIDisclosure_{t})", "\\overline{AIPatent}_{t}", "\\sigma(AIPatent_{t})"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("AIW", "it"),
                    math_text(" = "),
                    math_frac(
                        [
                            math_sub("AIDisclosure", "it"),
                            math_text(" - "),
                            math_sub_expr(math_overbar(math_rich_text("AIDisclosure", italic=False)), "t"),
                        ],
                        [
                            math_text("σ"),
                            math_group([math_sub("AIDisclosure", "t")]),
                        ],
                    ),
                    math_text(" - "),
                ],
                [
                    math_frac(
                        [
                            math_sub("AIPatent", "it"),
                            math_text(" - "),
                            math_sub_expr(math_overbar(math_rich_text("AIPatent", italic=False)), "t"),
                        ],
                        [
                            math_text("σ"),
                            math_group([math_sub("AIPatent", "t")]),
                        ],
                    ),
                ],
            ],
            eq_no=4,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "AIW_{it}=", "z(AIDisclosure_{it})", "z(AIPatent_{it})"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("AIW", "it"),
                math_text(" = z"),
                math_group([math_sub("AIDisclosure", "it")]),
                math_text(" - z"),
                math_group([math_sub("AIPatent", "it")]),
            ]],
            eq_no=3,
            lang=lang,
        )
        return True
    if text.startswith("AIW_it = z(AIDisclosure_it) - z(AIPatent_it)"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("AIW", "it"),
                math_text(" = z"),
                math_group([math_sub("AIDisclosure", "it")]),
                math_text(" - z"),
                math_group([math_sub("AIPatent", "it")]),
            ]],
            eq_no=4,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "Resilience_{it}=", "AIW_{it}", "Controls_{it}", "μ_i", "λ_t"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("Resilience", "it"),
                    math_text(" = "),
                    math_sub("α", "0"),
                    math_text(" + "),
                    math_sub("α", "1"),
                    math_text(" "),
                    math_sub("AIW", "it"),
                    math_text(" + "),
                    math_sub("α", "2"),
                    math_text(" "),
                    math_subsup("AIW", "it", "2"),
                    math_text(" + Σ "),
                    math_sub("β", "k"),
                    math_text(" "),
                    math_sub("Controls", "it"),
                ],
                [
                    math_text("+ "),
                    math_sub("μ", "i"),
                    math_text(" + "),
                    math_sub("λ", "t"),
                    math_text(" + "),
                    math_sub("ε", "it"),
                ],
            ],
            eq_no=4,
            lang=lang,
        )
        return True
    if text.startswith("Resilience_it = alpha_0 + alpha_1 AIW_it"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Resilience", "it"),
                math_text(" = "),
                math_sub("α", "0"),
                math_text(" + "),
                math_sub("α", "1"),
                math_text(" "),
                math_sub("AIW", "it"),
                math_text(" + "),
                math_sub("α", "2"),
                math_text(" "),
                math_subsup("AIW", "it", "2"),
                math_text(" + Σ "),
                math_sub("β", "k"),
                math_text(" "),
                math_sub("Controls", "it"),
                math_text(" + "),
                math_sub("μ", "i"),
                math_text(" + "),
                math_sub("λ", "t"),
                math_text(" + "),
                math_sub("ε", "it"),
            ]],
            eq_no=5,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "Channel_{it}=", "AIW_{it}", "Controls_{it}", "ν_{it}"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("Channel", "it"),
                    math_text(" = "),
                    math_sub("γ", "0"),
                    math_text(" + "),
                    math_sub("γ", "1"),
                    math_text(" "),
                    math_sub("AIW", "it"),
                    math_text(" + "),
                    math_sub("γ", "2"),
                    math_text(" "),
                    math_subsup("AIW", "it", "2"),
                    math_text(" + Σ "),
                    math_sub("δ", "k"),
                    math_text(" "),
                    math_sub("Controls", "it"),
                ],
                [
                    math_text("+ "),
                    math_sub("μ", "i"),
                    math_text(" + "),
                    math_sub("λ", "t"),
                    math_text(" + "),
                    math_sub("ν", "it"),
                ],
            ],
            eq_no=5,
            lang=lang,
        )
        return True
    if text.startswith("Channel_it = gamma_0 + gamma_1 AIW_it"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Channel", "it"),
                math_text(" = "),
                math_sub("γ", "0"),
                math_text(" + "),
                math_sub("γ", "1"),
                math_text(" "),
                math_sub("AIW", "it"),
                math_text(" + "),
                math_sub("γ", "2"),
                math_text(" "),
                math_subsup("AIW", "it", "2"),
                math_text(" + Σ "),
                math_sub("δ", "k"),
                math_text(" "),
                math_sub("Controls", "it"),
                math_text(" + "),
                math_sub("μ", "i"),
                math_text(" + "),
                math_sub("λ", "t"),
                math_text(" + "),
                math_sub("ν", "it"),
            ]],
            eq_no=6,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "Resilience_{it}=", "Channel_{it}", "ξ_{it}"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("Resilience", "it"),
                    math_text(" = "),
                    math_sub("θ", "0"),
                    math_text(" + "),
                    math_sub("θ", "1"),
                    math_text(" "),
                    math_sub("AIW", "it"),
                    math_text(" + "),
                    math_sub("θ", "2"),
                    math_text(" "),
                    math_subsup("AIW", "it", "2"),
                    math_text(" + "),
                    math_sub("θ", "3"),
                    math_text(" "),
                    math_sub("Channel", "it"),
                ],
                [
                    math_text("+ Σ "),
                    math_sub("φ", "k"),
                    math_text(" "),
                    math_sub("Controls", "it"),
                    math_text(" + "),
                    math_sub("μ", "i"),
                    math_text(" + "),
                    math_sub("λ", "t"),
                    math_text(" + "),
                    math_sub("ξ", "it"),
                ],
            ],
            eq_no=6,
            lang=lang,
        )
        return True
    if text.startswith("Resilience_it = theta_0 + theta_1 AIW_it"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Resilience", "it"),
                math_text(" = "),
                math_sub("θ", "0"),
                math_text(" + "),
                math_sub("θ", "1"),
                math_text(" "),
                math_sub("AIW", "it"),
                math_text(" + "),
                math_sub("θ", "2"),
                math_text(" "),
                math_subsup("AIW", "it", "2"),
                math_text(" + "),
                math_sub("θ", "3"),
                math_text(" "),
                math_sub("Channel", "it"),
                math_text(" + Σ "),
                math_sub("φ", "k"),
                math_text(" "),
                math_sub("Controls", "it"),
                math_text(" + "),
                math_sub("μ", "i"),
                math_text(" + "),
                math_sub("λ", "t"),
                math_text(" + "),
                math_sub("ξ", "it"),
            ]],
            eq_no=7,
            lang=lang,
        )
        return True
    if _has_formula_token(normalized, "Outcome_{it}=", "AnalystAttention_{it}", "κ_k", "ω_{it}"):
        replace_with_omml_block(
            paragraph,
            [
                [
                    math_sub("Outcome", "it"),
                    math_text(" = "),
                    math_sub("ρ", "0"),
                    math_text(" + "),
                    math_sub("ρ", "1"),
                    math_text(" "),
                    math_sub("AIW", "it"),
                    math_text(" + "),
                    math_sub("ρ", "2"),
                    math_text(" "),
                    math_subsup("AIW", "it", "2"),
                    math_text(" + "),
                    math_sub("ρ", "3"),
                    math_text(" "),
                    math_sub("AnalystAttention", "it"),
                ],
                [
                    math_text("+ "),
                    math_sub("ρ", "4"),
                    math_text(" "),
                    math_group([math_sub("AIW", "it"), math_text(" × "), math_sub("AnalystAttention", "it")]),
                    math_text(" + "),
                    math_sub("ρ", "5"),
                    math_text(" "),
                    math_group([math_subsup("AIW", "it", "2"), math_text(" × "), math_sub("AnalystAttention", "it")]),
                ],
                [
                    math_text("+ Σ "),
                    math_sub("κ", "k"),
                    math_text(" "),
                    math_sub("Controls", "it"),
                    math_text(" + "),
                    math_sub("μ", "i"),
                    math_text(" + "),
                    math_sub("λ", "t"),
                    math_text(" + "),
                    math_sub("ω", "it"),
                ],
            ],
            eq_no=7,
            lang=lang,
        )
        return True
    if text.startswith("Outcome_it = rho_0 + rho_1 AIW_it"):
        replace_with_omml_block(
            paragraph,
            [[
                math_sub("Outcome", "it"),
                math_text(" = "),
                math_sub("ρ", "0"),
                math_text(" + "),
                math_sub("ρ", "1"),
                math_text(" "),
                math_sub("AIW", "it"),
                math_text(" + "),
                math_sub("ρ", "2"),
                math_text(" "),
                math_subsup("AIW", "it", "2"),
                math_text(" + "),
                math_sub("ρ", "3"),
                math_text(" "),
                math_sub("AnalystAttention", "it"),
                math_text(" + "),
                math_sub("ρ", "4"),
                math_text(" "),
                math_group([math_sub("AIW", "it"), math_text(" × "), math_sub("AnalystAttention", "it")]),
                math_text(" + "),
                math_sub("ρ", "5"),
                math_text(" "),
                math_group([math_subsup("AIW", "it", "2"), math_text(" × "), math_sub("AnalystAttention", "it")]),
                math_text(" + Σ "),
                math_sub("κ", "k"),
                math_text(" "),
                math_sub("Controls", "it"),
                math_text(" + "),
                math_sub("μ", "i"),
                math_text(" + "),
                math_sub("λ", "t"),
                math_text(" + "),
                math_sub("ω", "it"),
            ]],
            eq_no=8,
            lang=lang,
        )
        return True
    return False


def rebuild_explanation_paragraph(paragraph, lang: str):
    east_asia = "宋体" if lang == "cn" else "Times New Roman"
    latin = "Times New Roman"
    size_pt = 10.5 if lang == "cn" else 11
    text = paragraph.text.strip()
    if lang == "cn":
        if text.startswith("其中，`W_kl` 表示节点"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "kl", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示节点 k 与节点 l 的共现强度，I(k in p) 表示专利 p 是否包含技术节点 k。这一矩阵并不只记录专利数量，而是刻画 AI 技术知识在全球专利中的组合关系：两个 IPC 节点共同出现在 AI 专利中的频次越高，说明相应技术方向之间的知识关联越紧密。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，`PR_k` 表示技术节点"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "k", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示技术节点 k 的 PageRank 权重，N 为网络中的技术节点数量，d 为阻尼系数，", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "lk", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示节点 l 指向节点 k 的共现权重，", east_asia, latin, size_pt)
            append_symbol(paragraph, "Σ", "m", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " ", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "lm", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示节点 l 的加权外向连接总强度。为反映技术网络结构随时间演化的特征，本文优先采用当年全球 AI 专利网络计算得到的节点权重；当某些年份节点信息不足时，采用全样本期网络权重进行补充。由此得到的权重可以区分不同 AI 专利所处技术位置的质量差异。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，`PR_kt` 为技术节点"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 为技术节点 k 在年份 t 的网络权重。该指标的含义是：企业并非因为拥有更多 AI 专利就自动获得更高技术行动得分，而是当其 AI 专利集中在全球 AI 知识网络中更核心、更具结构影响力的技术节点时，才获得更高的质量调整后 AI 技术积累得分。因此，AI Patent 衡量的是企业可验证 AI 技术资产的网络加权存量，而不是未经区分的专利件数。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，`Return_{im}` 表示企业") or text.startswith("其中，Return_{im} 表示企业"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "Return", "im", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 内第 m 月的个股收益率。该指标越小，说明企业在不确定环境中的股价波动越低、风险暴露越平缓，其吸收冲击和维持稳定的能力越强。", east_asia, latin, size_pt)
            return True
        if (
            text.startswith("其中，`w_1` 和 `w_2` 为熵权法确定的权重")
            or text.startswith("其中，w1和 w2为熵权法确定的权重")
            or text.startswith("其中，w1 和 w2 为熵权法确定的权重")
            or text.startswith("其中，w1和w2为熵权法确定的权重")
        ):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "w", "1", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 和 ", east_asia, latin, size_pt)
            append_symbol(paragraph, "w", "2", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 为熵权法确定的权重。因此，CR 取值越大，表明企业同时具备更强的短期稳定能力和中期恢复增长能力，即整体韧性水平越高。稳健性部分将进一步采用替代性韧性指标进行检验。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，`AIDisclosure_{it}` 表示企业") or text.startswith("其中，AIDisclosure_{it} 表示企业"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIDisclosure", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 的人工智能披露强度；", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIPatent", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 的人工智能专利得分。相应变量的年度横截面均值与年度横截面标准差分别用于完成同年样本内标准化。AIW 数值越大，说明企业人工智能叙事越明显领先于其实质性人工智能技术积累，人工智能洗白程度越高；反之，则说明企业人工智能技术积累相对更充分而公开叙事较为克制。为识别潜在的非线性关系，本文进一步构造平方项 ", east_asia, latin, size_pt)
            append_run(paragraph, "AIW", east_asia, latin, size_pt)
            append_run(paragraph, "²", east_asia, latin, size_pt)
            append_run(paragraph, "。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，PR_{kt} 表示技术节点"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示技术节点 k 在年份 t 的 PageRank 权重；M(k) 表示与节点 k 存在指向或连接关系的节点集合；", east_asia, latin, size_pt)
            append_symbol(paragraph, "L", "jt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示节点 j 在年份 t 的外向连接数量；d 为阻尼系数，通常设定为0.85。", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 越高，说明该技术节点在全球人工智能知识网络中越处于核心位置，其对应专利所代表的知识影响力和技术关联价值越高。为兼顾技术演进与样本覆盖，本文优先使用年度 PageRank 权重；若个别年份节点权重无法观测，则使用全样本期网络权重进行补充。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，PRkt 表示技术节点"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示技术节点 k 在年份 t 的 PageRank 权重；M(k) 表示与节点 k 存在指向或连接关系的节点集合；", east_asia, latin, size_pt)
            append_symbol(paragraph, "L", "jt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示节点 j 在年份 t 的外向连接数量；d 为阻尼系数，通常设定为0.85。", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 越高，说明该技术节点在全球人工智能知识网络中越处于核心位置，其对应专利所代表的知识影响力和技术关联价值越高。为兼顾技术演进与样本覆盖，本文优先使用年度 PageRank 权重；若个别年份节点权重无法观测，则使用全样本期网络权重进行补充。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，K_{it} 表示企业"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "K", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 所涉及的人工智能技术节点集合；", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示技术节点 k 在年份 t 的 PageRank 权重；", east_asia, latin, size_pt)
            append_symbol(paragraph, "N", "ikt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 于节点 k 上的人工智能专利申请数量。该指标不仅反映企业人工智能专利的数量，还进一步考虑企业所进入技术领域在全球 AI 知识网络中的结构位置、知识关联程度与技术重要性，因此能够比原始专利计数更准确地刻画企业真实人工智能技术能力。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，z(AIDisclosure_it) 和 z(AIPatent_it)"):
            clear_paragraph(paragraph)
            append_run(paragraph, "其中，z(", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIDisclosure", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, ") 和 z(", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIPatent", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, ") 分别表示 AI Disclosure 与 AI Patent 在年份 t 内的横截面标准化值。AIW 数值越大，表示企业人工智能叙事越明显领先于其实质性人工智能技术积累，人工智能洗白程度越高。需要说明的是，这一指标本质上刻画的是“AI 叙事 - 能力基础偏离”而非对企业主观意图的直接观测，因此更接近 AI washing 的可操作代理变量，而不是对“误导行为”本身的完全识别。之所以仍将其用于度量 AI washing，是因为在上市公司情境下，外部受众最先观察到的正是公开叙事，而专利积累虽然不能穷尽企业全部 AI 能力，却能够较稳定地反映其可验证技术行动基础。当企业 AI 叙事持续显著快于这一可验证基础时，外部受众面对的正是本文所关注的“高叙事可见性、低能力可即时验证性”的偏离状态。为识别潜在的非线性关系，本文进一步构造平方项 ", east_asia, latin, size_pt)
            append_run(paragraph, "AIW", east_asia, latin, size_pt)
            append_run(paragraph, "²", east_asia, latin, size_pt)
            append_run(paragraph, "。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，Resilience_it 表示企业 i 在年份 t 的企业韧性"):
            reset_paragraph_runs(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "Resilience", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 表示企业 i 在年份 t 的企业韧性；", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIW", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 为人工智能洗白程度；", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIW", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, "² 为其平方项；", east_asia, latin, size_pt)
            append_symbol(paragraph, "Controls", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 为控制变量集合；", east_asia, latin, size_pt)
            append_symbol(paragraph, "μ", "i", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " 为企业固定效应；", east_asia, latin, size_pt)
            append_symbol(paragraph, "λ", "t", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " 为年份固定效应；", east_asia, latin, size_pt)
            append_symbol(paragraph, "ε", "it", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " 为随机扰动项。若 ", east_asia, latin, size_pt)
            append_symbol(paragraph, "α", "1", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " > 0 且 ", east_asia, latin, size_pt)
            append_symbol(paragraph, "α", "2", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " < 0，则表明人工智能洗白与企业韧性之间存在显著的倒 U 型关系。", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，Channel_it 分别表示 Trade Credit 和 Agency Cost"):
            reset_paragraph_runs(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "Channel", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 分别表示 Trade Credit 和 Agency Cost。在此基础上，再将中介变量纳入企业韧性回归：", east_asia, latin, size_pt)
            return True
        if text.startswith("其中，Outcome_it 分别表示 Resilience、Trade Credit 和 Agency Cost"):
            reset_paragraph_runs(paragraph)
            append_run(paragraph, "其中，", east_asia, latin, size_pt)
            append_symbol(paragraph, "Outcome", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " 分别表示 Resilience、Trade Credit 和 Agency Cost。为降低交互项带来的多重共线性问题，本文在构造交互项前对 AIW 与 Analyst Attention 进行中心化处理。该模型不仅能够检验分析师关注度是否改变人工智能洗白的边际影响方向与强度，也能够识别其是否改变倒 U 型关系的拐点位置。", east_asia, latin, size_pt)
            return True
    else:
        if text.startswith("where `W_kl` is the co-occurrence strength"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "kl", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the co-occurrence strength between nodes k and l, and I(k in p) indicates whether patent p contains technology node k. The matrix therefore captures how AI-related technological knowledge is combined across patent classes. A higher co-occurrence frequency indicates a tighter knowledge association between the corresponding technology areas.", east_asia, latin, size_pt)
            return True
        if text.startswith("where `PR_k` is the PageRank weight"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "k", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the PageRank weight of technology node k, N is the number of nodes in the network, d is the damping factor, ", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "lk", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the weighted co-occurrence link from node l to node k, and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Σ", "m", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " ", east_asia, latin, size_pt)
            append_symbol(paragraph, "W", "lm", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the total weighted outward connection strength of node l. To reflect the evolution of the technology network over time, we use year-specific PageRank weights whenever available and supplement them with full-sample network weights when node information is insufficient in a given year. These weights distinguish AI patents by the structural importance of the technological knowledge they embody.", east_asia, latin, size_pt)
            return True
        if text.startswith("where `PR_kt` is the network weight"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the network weight of node k in year t. This score means that a firm does not receive a higher AI-action score simply because it files more AI patents. It receives a higher score when its AI patents are concentrated in more central and structurally influential positions within the global AI knowledge network. The resulting AI Patent variable therefore measures a firm's quality-adjusted stock of verifiable AI technological assets rather than an undifferentiated patent count.", east_asia, latin, size_pt)
            return True
        if text.startswith("where `Return_{im}` denotes the stock return") or text.startswith("where Return_{im} denotes the stock return"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Return", "im", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " denotes the stock return of firm i in month m of year t. Lower values indicate lower volatility and stronger shock-absorption capability.", east_asia, latin, size_pt)
            return True
        if text.startswith("where `w_1` and `w_2` are the entropy-based weights"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "w", "1", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "w", "2", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " are the entropy-based weights. Higher values of CR therefore indicate stronger overall resilience, reflecting both short-term stability and medium-term recovery growth. Robustness checks use an alternative resilience measure.", east_asia, latin, size_pt)
            return True
        if (
            text.startswith("where `AIDisclosure_{it}` denotes firm `i`'s AI disclosure intensity")
            or text.startswith("where AIDisclosure_{it} denotes firm i's AI disclosure intensity")
            or text.startswith("where AIDisclosure_{it} denotes firm i’s AI disclosure intensity")
        ):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIDisclosure", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " denotes firm i's AI disclosure intensity in year t, and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIPatent", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " denotes the firm's AI patent score in year t. The yearly cross-sectional means and standard deviations of the two components are used to standardize them within the same year. Higher values indicate that AI-related narratives run further ahead of substantive AI technological accumulation. We further construct the squared term ", east_asia, latin, size_pt)
            append_run(paragraph, "AIW", east_asia, latin, size_pt)
            append_run(paragraph, "²", east_asia, latin, size_pt)
            append_run(paragraph, " to test the nonlinear relationship.", east_asia, latin, size_pt)
            return True
        if text.startswith("where PR_{kt} is the PageRank weight"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the PageRank weight of technology node k in year t, M(k) is the set of nodes linked to node k, ", east_asia, latin, size_pt)
            append_symbol(paragraph, "L", "jt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the number of outgoing links of node j in year t, and d is the damping factor, conventionally set at 0.85. A higher ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " indicates that the node occupies a more central position in the global AI knowledge network.", east_asia, latin, size_pt)
            return True
        if text.startswith("where PRkt is the PageRank weight"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the PageRank weight of technology node k in year t, M(k) is the set of nodes linked to node k, ", east_asia, latin, size_pt)
            append_symbol(paragraph, "L", "jt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the number of outgoing links of node j in year t, and d is the damping factor, conventionally set at 0.85. A higher ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " indicates that the node occupies a more central position in the global AI knowledge network.", east_asia, latin, size_pt)
            return True
        if text.startswith("where K_{it} is the set of AI technology nodes"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "K", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the set of AI technology nodes covered by firm i in year t, ", east_asia, latin, size_pt)
            append_symbol(paragraph, "PR", "kt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the PageRank weight of node k, and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "N", "ikt", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the number of AI patent applications of firm i on node k in year t. This measure captures both patent quantity and the structural importance of the technology fields in which the firm is active.", east_asia, latin, size_pt)
            return True
        if text.startswith("where z(AIDisclosure_it) and z(AIPatent_it)"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where z(", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIDisclosure", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, ") and z(", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIPatent", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, ") denote the within-year cross-sectional standardized values of AI Disclosure and AI Patent, respectively. Higher values of AIW indicate that a firm's AI-oriented narrative is more pronounced relative to its substantive AI technological accumulation. Importantly, this measure captures rhetoric-capability divergence rather than managerial intent directly. It is therefore best interpreted as an operational proxy for AI washing rather than a complete observation of deceptive behavior itself. We nonetheless use it as our empirical measure because, in listed-firm settings, outside audiences first observe public rhetoric, whereas patent accumulation, while not exhausting all AI capability, provides a relatively stable indicator of verifiable technological action. When AI rhetoric persistently and substantially outruns that verifiable base, outside audiences face precisely the high-visibility, low-immediate-verifiability condition that defines the phenomenon of interest here. To test for nonlinearity, we further construct the squared term ", east_asia, latin, size_pt)
            append_run(paragraph, "AIW", east_asia, latin, size_pt)
            append_run(paragraph, "²", east_asia, latin, size_pt)
            append_run(paragraph, ".", east_asia, latin, size_pt)
            return True
        if text.startswith("where Resilience_it denotes corporate resilience for firm i in year t"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Resilience", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " denotes corporate resilience for firm i in year t; ", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIW", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is the degree of AI washing; ", east_asia, latin, size_pt)
            append_symbol(paragraph, "AIW", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, "² is its squared term; ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Controls", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is a vector of controls; ", east_asia, latin, size_pt)
            append_symbol(paragraph, "μ", "i", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " denotes firm fixed effects; ", east_asia, latin, size_pt)
            append_symbol(paragraph, "λ", "t", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " denotes year fixed effects; and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "ε", "it", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " is the disturbance term. An inverted U-shaped relationship is supported when ", east_asia, latin, size_pt)
            append_symbol(paragraph, "α", "1", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " > 0 and ", east_asia, latin, size_pt)
            append_symbol(paragraph, "α", "2", east_asia, latin, size_pt, greek=True)
            append_run(paragraph, " < 0.", east_asia, latin, size_pt)
            return True
        if text.startswith("where Channel_it is alternately Trade Credit and Agency Cost"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Channel", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is alternately Trade Credit and Agency Cost. We then add the mediator to the resilience equation:", east_asia, latin, size_pt)
            return True
        if text.startswith("where Outcome_it is alternately Resilience, Trade Credit, and Agency Cost"):
            clear_paragraph(paragraph)
            append_run(paragraph, "where ", east_asia, latin, size_pt)
            append_symbol(paragraph, "Outcome", "it", east_asia, latin, size_pt, italic=False)
            append_run(paragraph, " is alternately Resilience, Trade Credit, and Agency Cost. To reduce multicollinearity in the interaction terms, we center AIW and Analyst Attention before constructing the interaction terms. This specification allows us to test not only whether analyst attention changes the marginal effect of AI washing, but also whether it shifts the turning point of the nonlinear relationship.", east_asia, latin, size_pt)
            return True
    return False


def format_paragraphs(doc: Document, lang: str):
    east_asia = "宋体" if lang == "cn" else "Times New Roman"
    latin = "Times New Roman"
    body_size = 10.5 if lang == "cn" else 11
    for paragraph in list(doc.paragraphs):
        if rebuild_equation_paragraph(paragraph, lang):
            continue
        if rebuild_explanation_paragraph(paragraph, lang):
            set_paragraph_format(paragraph, lang, in_table=False)
            continue
        if rebuild_inline_symbol_paragraph(paragraph, lang):
            set_paragraph_format(paragraph, lang, in_table=False)
            continue
        if rebuild_collapsed_symbol_paragraph(paragraph, lang):
            set_paragraph_format(paragraph, lang, in_table=False)
            continue
        set_paragraph_format(paragraph, lang, in_table=False)
        for run in paragraph.runs:
            if paragraph_has_drawing(paragraph):
                continue
            if lang == "cn" and run.text:
                run.text = normalize_cn_text(run.text)
            size = body_size
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name == "Heading 1":
                size = 15
            elif style_name == "Heading 2":
                size = 13
            elif style_name == "Heading 3":
                size = 12
            elif is_caption_text(paragraph.text.strip()):
                size = 10.5
            set_run_fonts(run, east_asia, latin, size)


def _paragraph_text_xml(p) -> str:
    return "".join(t.text or "" for t in p.findall(f".//{{{W_NS}}}t"))


def _is_equation_paragraph_xml(p) -> bool:
    return bool(p.findall(f".//{{{M_NS}}}oMathPara") or p.findall(f".//{{{M_NS}}}oMath"))


def _parse_equation_number(text: str) -> int | None:
    cleaned = text.strip().replace("(", "（").replace(")", "）")
    if cleaned.startswith("（") and cleaned.endswith("）"):
        inner = cleaned[1:-1].strip()
        if inner.isdigit():
            return int(inner)
    return None


def consolidate_existing_equation_numbers(doc: Document):
    """Merge equation paragraphs plus standalone number paragraphs into native OMML rows."""
    body = doc._element.body
    children = list(body)
    idx = 0
    while idx < len(children) - 1:
        current = children[idx]
        nxt = children[idx + 1]
        if current.tag != qn("w:p") or nxt.tag != qn("w:p"):
            idx += 1
            continue
        if not _is_equation_paragraph_xml(current):
            idx += 1
            continue
        eq_no = _parse_equation_number(_paragraph_text_xml(nxt))
        if eq_no is None:
            idx += 1
            continue
        lines = _extract_math_lines_from_paragraph_xml(current)
        if not lines:
            idx += 1
            continue
        new_p = build_equation_numbered_paragraph(lines, eq_no, lang="cn")
        body.insert(body.index(current), new_p)
        body.remove(current)
        body.remove(nxt)
        children = list(body)
        idx += 1


def remove_orphan_equation_number_paragraphs(doc: Document):
    body = doc._element.body
    children = list(body)
    for idx, node in enumerate(list(children)):
        if node.tag != qn("w:p"):
            continue
        eq_no = _parse_equation_number(_paragraph_text_xml(node))
        if eq_no is None:
            continue
        prev_node = children[idx - 1] if idx > 0 else None
        if prev_node is None:
            continue
        if prev_node.tag == qn("w:p") and _is_equation_paragraph_xml(prev_node):
            body.remove(node)
        elif prev_node.tag == qn("w:tbl") and ("<m:oMath" in prev_node.xml or "<m:oMathPara" in prev_node.xml):
            body.remove(node)


def renumber_equation_tables(doc: Document):
    body = doc._element.body
    eq_no = 1
    for child in list(body):
        if child.tag != qn("w:tbl"):
            continue
        xml = child.xml
        if "<m:oMath" not in xml and "<m:oMathPara" not in xml:
            continue
        cells = child.findall(f".//{{{W_NS}}}tc")
        if len(cells) < 2:
            continue
        text_nodes = cells[-1].findall(f".//{{{W_NS}}}t")
        if not text_nodes:
            continue
        text_nodes[0].text = f"（{eq_no}）"
        for node in text_nodes[1:]:
            node.text = ""
        eq_no += 1


def force_caption_alignment(doc: Document):
    for p in doc._element.body.iter(qn("w:p")):
        text = _paragraph_text_xml(p).strip()
        if not is_caption_text(text):
            continue
        p_pr = p.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = _w_el("pPr")
            p.insert(0, p_pr)
        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = _w_el("jc")
            p_pr.append(jc)
        jc.set(qn("w:val"), "center")
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = _w_el("ind")
            p_pr.append(ind)
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:right"), "0")
        ind.set(qn("w:firstLine"), "0")


def normalize_plain_chinese_paragraphs(doc: Document):
    """Normalize visible Chinese text when it is safe to collapse run text.

    This deliberately skips paragraphs containing fields, drawings, or OMML so
    Zotero live citations, images, and equations are not touched.
    """
    for p in doc._element.body.iter(qn("w:p")):
        if (
            p.findall(f".//{{{W_NS}}}fldChar")
            or p.findall(f".//{{{W_NS}}}instrText")
            or p.findall(f".//{{{W_NS}}}drawing")
            or p.findall(f".//{{{M_NS}}}oMath")
            or p.findall(f".//{{{M_NS}}}oMathPara")
            or p.findall(f".//{{{W_NS}}}vertAlign")
        ):
            continue
        text_nodes = p.findall(f".//{{{W_NS}}}t")
        if not text_nodes:
            continue
        original = "".join(node.text or "" for node in text_nodes)
        normalized = normalize_caption_text(original) if is_caption_text(original) else normalize_cn_text(original)
        if normalized == original:
            continue
        text_nodes[0].text = normalized
        for node in text_nodes[1:]:
            node.text = ""


def audit_inline_symbol_delivery(doc: Document):
    failures = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_has_fields(paragraph) or paragraph_has_drawing(paragraph):
            continue
        visible = "".join(run.text for run in paragraph.runs)
        if not visible:
            continue
        subscript_runs = {run.text for run in paragraph.runs if run.font.subscript}
        for token, expected_sub in INLINE_SUBSCRIPT_AUDIT_PATTERNS.items():
            if token in visible and expected_sub not in subscript_runs:
                failures.append((idx, token, visible[:160]))
                break
    if failures:
        preview = "; ".join(f"p{idx}:{token}" for idx, token, _ in failures[:8])
        raise RuntimeError(f"Inline symbol delivery audit failed: {preview}")


def audit_equation_paragraph_delivery(doc: Document):
    failures = []
    for idx, paragraph in enumerate(doc.paragraphs):
        visible = "".join(run.text for run in paragraph.runs).strip()
        has_math = _is_equation_paragraph_xml(paragraph._element)
        if _parse_equation_number(visible) is not None and not has_math:
            failures.append((idx, visible))
            continue
        if has_math and EQUATION_NO_ONLY_RE.fullmatch(visible):
            if "w:tab" not in paragraph._element.xml:
                failures.append((idx, visible))
    if failures:
        preview = "; ".join(f"p{idx}:{text}" for idx, text in failures[:8])
        raise RuntimeError(f"Equation numbering audit failed: {preview}")


def audit_equation_layout_tables(doc: Document, *, allow_fallback: bool = False):
    """Block table-based numbered equations unless an explicit fallback was requested."""
    if allow_fallback:
        return
    failures = []
    for idx, table in enumerate(doc.tables):
        xml = table._element.xml
        if ("<m:oMath" not in xml and "<m:oMathPara" not in xml) or _parse_equation_number(table_visible_text(table)) is None:
            continue
        failures.append(idx)
    if failures:
        preview = ", ".join(f"table{idx}" for idx in failures[:8])
        raise RuntimeError(
            "Equation layout audit failed: numbered equations must use native equation paragraphs "
            f"with center/right tab stops, not table containers ({preview}). "
            "Use --allow-equation-table-fallback only for a documented recovery exception."
        )


def audit_table_cell_paragraph_geometry(doc: Document):
    """Fail if ordinary academic tables still inherit body paragraph geometry."""
    failures = []
    for table_idx, table in enumerate(doc.tables):
        if is_equation_table(table):
            continue
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.tcPr
                v_align = tc_pr.find(qn("w:vAlign")) if tc_pr is not None else None
                if v_align is None or v_align.get(qn("w:val")) != "center":
                    failures.append(f"table{table_idx} r{row_idx} c{cell_idx}:vAlign")
                    continue
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    p_pr = paragraph._element.find(qn("w:pPr"))
                    ind = p_pr.find(qn("w:ind")) if p_pr is not None else None
                    jc = p_pr.find(qn("w:jc")) if p_pr is not None else None
                    if ind is None:
                        failures.append(f"table{table_idx} r{row_idx} c{cell_idx} p{p_idx}:missing-ind")
                        continue
                    required_zero_attrs = ("left", "right", "firstLine", "firstLineChars")
                    bad_attrs = [attr for attr in required_zero_attrs if ind.get(qn(f"w:{attr}")) != "0"]
                    if bad_attrs:
                        failures.append(f"table{table_idx} r{row_idx} c{cell_idx} p{p_idx}:bad-{','.join(bad_attrs)}")
                        continue
                    if any(ind.get(qn(f"w:{attr}")) is not None for attr in ("hanging", "hangingChars")):
                        failures.append(f"table{table_idx} r{row_idx} c{cell_idx} p{p_idx}:hanging-indent")
                        continue
                    if jc is None or jc.get(qn("w:val")) != "center":
                        failures.append(f"table{table_idx} r{row_idx} c{cell_idx} p{p_idx}:jc")
                        continue
    if failures:
        preview = "; ".join(failures[:12])
        raise RuntimeError(
            "Table-cell paragraph geometry audit failed: ordinary academic table cells must "
            f"have explicit zero indentation and centered alignment ({preview})."
        )


def apply_page_break_rules(doc: Document, lang: str):
    if lang == "cn":
        break_before_titles = {"摘要", "引言", "理论基础与研究假设", "研究设计", "实证结果分析", "进一步分析", "讨论", "结论", "参考文献"}
    else:
        break_before_titles = {"Abstract", "Introduction", "Theory and Hypotheses", "Research Design", "Empirical Results", "Additional Analyses", "Discussion", "Conclusion", "References", "参考文献"}
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.page_break_before = paragraph.text.strip() in break_before_titles


def validate_docx(path: Path, citation_policy: str = "strict"):
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    if "�" in xml or "????" in xml:
        raise RuntimeError(f"{path} contains garbling markers")
    if "ADDIN ZOTERO_ITEM" not in xml and "CSL_CITATION" not in xml and citation_policy == "strict":
        raise RuntimeError(f"{path} lost live citation fields")
    if "ADDIN ZOTERO_ITEM" not in xml and "CSL_CITATION" not in xml and citation_policy == "warn":
        print(f"WARNING: {path} has no live citation fields; layout finalization completed as a draft/recovery pass only.")
    for old_name in ("CR_it", "Y_it", "Mediator_it", "ROA1", "Balance1", "Mshare", "Occupy"):
        if old_name in xml:
            raise RuntimeError(f"{path} still contains deprecated variable token: {old_name}")
    for raw_formula in (
        "Stability_{it}",
        "Return_{im}",
        "Growth_{it}",
        "Revenue_{it}",
        "Revenue_{i,t-3}",
        "CR_{it}",
        "w_1",
        "w_2",
        "AIDisclosure_{it}",
        "AIPatent_{it}",
        "\\overline{AIDisclosure}_{t}",
        "\\overline{AIPatent}_{t}",
        "\\sigma(AIDisclosure_{t})",
        "\\sigma(AIPatent_{t})",
    ):
        if raw_formula in xml:
            raise RuntimeError(f"{path} still contains raw formula residue: {raw_formula}")
    if "<wp:anchor" in xml:
        raise RuntimeError(f"{path} contains floating anchor images")


def main():
    parser = argparse.ArgumentParser(description="Finalize bilingual submission DOCX with citation-safe post-processing.")
    parser.add_argument("--input-docx", required=True)
    parser.add_argument("--output-docx", required=True)
    parser.add_argument("--lang", choices=["cn", "en"], required=True)
    parser.add_argument("--mode", default="journal_submission")
    parser.add_argument(
        "--citation-policy",
        choices=["strict", "warn", "off"],
        default="strict",
        help="Use strict for formal delivery, warn for disaster-recovery or working-draft layout repair, and off only for non-citation-managed documents.",
    )
    parser.add_argument("--reference-style", default=None)
    parser.add_argument("--template-profile", default=None)
    parser.add_argument(
        "--allow-equation-table-fallback",
        action="store_true",
        help=(
            "Allow table-based numbered equation containers only for a documented recovery "
            "exception. Default formal delivery requires native equation paragraphs with "
            "center/right tab-stop numbering."
        ),
    )
    args = parser.parse_args()
    if args.mode != "journal_submission":
        raise ValueError("Only --mode journal_submission is currently supported")
    doc = Document(args.input_docx)
    set_style_fonts(doc, args.lang)
    format_paragraphs(doc, args.lang)
    normalize_existing_omml_numbered_paragraphs(doc, args.lang)
    consolidate_existing_equation_numbers(doc)
    normalize_display_equation_paragraphs(doc, args.lang)
    remove_orphan_equation_number_paragraphs(doc)
    if args.lang == "cn":
        normalize_plain_chinese_paragraphs(doc)
    force_caption_alignment(doc)
    apply_page_break_rules(doc, args.lang)
    format_tables(doc, args.lang)
    audit_inline_symbol_delivery(doc)
    audit_equation_paragraph_delivery(doc)
    audit_equation_layout_tables(doc, allow_fallback=args.allow_equation_table_fallback)
    audit_table_cell_paragraph_geometry(doc)
    out_path = Path(args.output_docx)
    doc.save(out_path)
    repair_equation_number_tabs(out_path, out_path)
    validate_docx(out_path, citation_policy=args.citation_policy)


if __name__ == "__main__":
    main()
