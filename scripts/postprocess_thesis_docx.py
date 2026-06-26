import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


EMU_PER_INCH = 914400


def set_run_fonts(run, east_asia="宋体", latin="Times New Roman", size_pt=12, bold=None, italic=None, subscript=False):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.subscript = subscript


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def append_text_run(paragraph, text, *, bold=None, italic=None, subscript=False, size_pt=12):
    run = paragraph.add_run(text)
    set_run_fonts(run, bold=bold, italic=italic, subscript=subscript, size_pt=size_pt)
    return run


def build_symbol(paragraph, base, sub=None, *, greek=False):
    run = paragraph.add_run(base)
    set_run_fonts(run, italic=not greek, size_pt=12)
    if sub:
        sub_run = paragraph.add_run(sub)
        set_run_fonts(sub_run, italic=not greek, subscript=True, size_pt=12)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        edge_el = tc_borders.find(qn(f"w:{edge}"))
        if edge_data is None:
            if edge_el is not None:
                tc_borders.remove(edge_el)
            continue
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tc_borders.append(edge_el)
        for key, value in edge_data.items():
            edge_el.set(qn(f"w:{key}"), str(value))


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)


def format_image_table(doc, table_index, image_indexes, target_width_in):
    table = doc.tables[table_index]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)
    remove_table_borders(table)

    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    col_count = len(table.columns)
    col_width = int((page_width / col_count) - 0.12 * EMU_PER_INCH)
    for col_idx in range(col_count):
        for cell in table.columns[col_idx].cells:
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=60, start=90, bottom=60, end=90)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    target_width = int(target_width_in * EMU_PER_INCH)
    shapes = list(doc.inline_shapes)
    for idx in image_indexes:
        if idx >= len(shapes):
            continue
        shape = shapes[idx]
        if shape.width > target_width:
            ratio = shape.height / shape.width
            shape.width = target_width
            shape.height = int(target_width * ratio)


def format_three_line_table(
    doc,
    table_index,
    col_widths_in,
    top_rule_pt,
    mid_rule_pt,
    bottom_rule_pt,
    header_font_pt,
    body_font_pt,
    header_line_pt,
    body_line_pt,
):
    table = doc.tables[table_index]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)

    col_widths = [int(v * EMU_PER_INCH) for v in col_widths_in]
    for col_idx, width in enumerate(col_widths):
        if col_idx >= len(table.columns):
            break
        for cell in table.columns[col_idx].cells:
            cell.width = width

    top = {"val": "single", "sz": str(int(top_rule_pt * 8)), "space": "0", "color": "000000"}
    mid = {"val": "single", "sz": str(int(mid_rule_pt * 8)), "space": "0", "color": "000000"}
    bottom = {"val": "single", "sz": str(int(bottom_rule_pt * 8)), "space": "0", "color": "000000"}
    none = {"val": "nil"}

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=36, start=72, bottom=36, end=72)
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.left_indent = Pt(0)
                pf.right_indent = Pt(0)
                pf.first_line_indent = Pt(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(header_line_pt if row_idx == 0 else body_line_pt)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (row_idx == 0 or col_idx < 3) else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_fonts(run, size_pt=header_font_pt if row_idx == 0 else body_font_pt, bold=(row_idx == 0))

            cell_top = top if row_idx == 0 else none
            cell_bottom = mid if row_idx == 0 else bottom if row_idx == len(table.rows) - 1 else none
            set_cell_borders(cell, top=cell_top, bottom=cell_bottom, left=none, right=none)


def replace_paragraph_if_matches(doc, startswith_text, builder):
    for p in doc.paragraphs:
        if p.text.startswith(startswith_text):
            clear_paragraph(p)
            builder(p)
            return True
    return False


def ai_patent_builder(paragraph):
    append_text_run(paragraph, "其中，")
    build_symbol(paragraph, "K", "it")
    append_text_run(paragraph, " 表示企业 ")
    build_symbol(paragraph, "i")
    append_text_run(paragraph, " 在年份 ")
    build_symbol(paragraph, "t")
    append_text_run(paragraph, " 所涉及的 AI 技术节点集合，")
    build_symbol(paragraph, "PR", "kt")
    append_text_run(paragraph, " 表示技术节点 ")
    build_symbol(paragraph, "k")
    append_text_run(paragraph, " 在年份 ")
    build_symbol(paragraph, "t")
    append_text_run(paragraph, " 的 PageRank 权重，")
    build_symbol(paragraph, "N", "ikt")
    append_text_run(paragraph, " 表示企业 ")
    build_symbol(paragraph, "i")
    append_text_run(paragraph, " 在年份 ")
    build_symbol(paragraph, "t")
    append_text_run(paragraph, " 于节点 ")
    build_symbol(paragraph, "k")
    append_text_run(paragraph, " 上的专利申请数量。")


def model_builder(paragraph):
    append_text_run(paragraph, "其中，企业固定效应记为 ")
    build_symbol(paragraph, "μ", "i", greek=True)
    append_text_run(paragraph, "，年份固定效应记为 ")
    build_symbol(paragraph, "λ", "t", greek=True)
    append_text_run(paragraph, "。若一次项系数 ")
    build_symbol(paragraph, "β", "1", greek=True)
    append_text_run(paragraph, " 为正、二次项系数 ")
    build_symbol(paragraph, "β", "2", greek=True)
    append_text_run(paragraph, " 为负，则与本文所提出的倒 U 型关系一致。")


def repair_formula_paragraphs(doc):
    replace_paragraph_if_matches(doc, "其中，(K_{it})", ai_patent_builder)
    replace_paragraph_if_matches(doc, "其中，企业固定效应记为", model_builder)
    replace_paragraph_if_matches(doc, "其中，(_i)", model_builder)


def main():
    parser = argparse.ArgumentParser(description="Post-process Chinese thesis DOCX for figures, tables, and formula explanations.")
    parser.add_argument("input_docx")
    parser.add_argument("output_docx")
    parser.add_argument("--image-table-index", type=int, default=0)
    parser.add_argument("--image-indexes", default="1,2")
    parser.add_argument("--image-target-width-in", type=float, default=2.82)
    parser.add_argument("--table-index", type=int, default=1)
    parser.add_argument("--table-col-widths-in", default="0.98,1.20,1.00,3.05")
    parser.add_argument("--top-rule-pt", type=float, default=1.5)
    parser.add_argument("--mid-rule-pt", type=float, default=0.5)
    parser.add_argument("--bottom-rule-pt", type=float, default=1.5)
    parser.add_argument("--header-font-pt", type=float, default=11)
    parser.add_argument("--body-font-pt", type=float, default=10.5)
    parser.add_argument("--header-line-pt", type=float, default=16)
    parser.add_argument("--body-line-pt", type=float, default=18)
    args = parser.parse_args()

    in_path = Path(args.input_docx)
    out_path = Path(args.output_docx)
    doc = Document(str(in_path))

    image_indexes = [int(x) for x in args.image_indexes.split(",") if x.strip()]
    col_widths_in = [float(x) for x in args.table_col_widths_in.split(",") if x.strip()]

    if len(doc.tables) > args.image_table_index:
        format_image_table(doc, args.image_table_index, image_indexes, args.image_target_width_in)
    if len(doc.tables) > args.table_index:
        format_three_line_table(
            doc,
            args.table_index,
            col_widths_in,
            args.top_rule_pt,
            args.mid_rule_pt,
            args.bottom_rule_pt,
            args.header_font_pt,
            args.body_font_pt,
            args.header_line_pt,
            args.body_line_pt,
        )
    repair_formula_paragraphs(doc)
    doc.save(str(out_path))


if __name__ == "__main__":
    main()
